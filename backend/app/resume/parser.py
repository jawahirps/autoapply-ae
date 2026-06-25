import re
import io
from pathlib import Path
from typing import Optional

try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    from pdfminer.high_level import extract_text as pdfminer_extract
    HAS_PDFMINER = True
except ImportError:
    HAS_PDFMINER = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def extract_text_from_pdf(file_path: str) -> str:
    text = ""
    if HAS_PDFMINER:
        try:
            text = pdfminer_extract(file_path)
            if text and text.strip():
                return text
        except Exception:
            pass
    if HAS_PYPDF2:
        try:
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() or ""
        except Exception as e:
            raise ValueError(f"Could not read PDF: {e}")
    return text


def extract_text_from_docx(file_path: str) -> str:
    if not HAS_DOCX:
        raise ValueError("python-docx not installed")
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)
    return "\n".join(paragraphs)


def extract_text(file_path: str) -> str:
    path = Path(file_path)
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def extract_email(text: str) -> Optional[str]:
    pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    matches = re.findall(pattern, text)
    return matches[0] if matches else None


def extract_phone(text: str) -> Optional[str]:
    patterns = [
        r'(?:\+971|00971|0)[\s\-]?(?:5[0-9]|2|3|4|6|7|9)[\s\-]?\d{3}[\s\-]?\d{4}',  # UAE
        r'\+?[\d\s\-\(\)]{10,15}',
    ]
    for pat in patterns:
        matches = re.findall(pat, text)
        if matches:
            return matches[0].strip()
    return None


def extract_name(text: str) -> Optional[str]:
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    # First non-empty, non-contact line is usually the name
    for line in lines[:5]:
        if (len(line.split()) >= 2
                and len(line) < 60
                and not any(c in line for c in ['@', 'http', ':', '|'])
                and not re.search(r'\d{4}', line)):
            return line
    return None


SKILLS_DB = [
    # Tech
    "python","java","javascript","typescript","react","angular","vue","node.js","nodejs",
    "django","flask","fastapi","spring","sql","mysql","postgresql","mongodb","redis",
    "docker","kubernetes","aws","azure","gcp","terraform","ansible","linux","git",
    "machine learning","deep learning","nlp","data science","tableau","power bi",
    "excel","word","powerpoint","sap","salesforce","jira","agile","scrum",
    # Business & Finance
    "accounting","finance","audit","tax","ifrs","gaap","quickbooks","xero",
    "financial modeling","valuation","risk management","compliance",
    # Marketing
    "digital marketing","seo","sem","google ads","facebook ads","content marketing",
    "social media","email marketing","crm","hubspot",
    # Engineering
    "autocad","solidworks","matlab","project management","pmp","prince2",
    "civil engineering","mechanical engineering","electrical engineering",
    # Languages
    "arabic","english","french","hindi","urdu","chinese","spanish",
    # Soft skills
    "leadership","communication","teamwork","problem solving","analytical",
]


def extract_skills(text: str) -> list:
    text_lower = text.lower()
    found = []
    for skill in SKILLS_DB:
        if re.search(r'\b' + re.escape(skill) + r'\b', text_lower):
            found.append(skill)
    return list(set(found))


def extract_experience(text: str) -> list:
    experience = []
    exp_section = ""
    lines = text.split('\n')
    in_exp = False
    exp_keywords = ['experience', 'employment', 'work history', 'professional background', 'career']
    stop_keywords = ['education', 'skills', 'projects', 'certifications', 'awards', 'references']

    for i, line in enumerate(lines):
        line_lower = line.lower().strip()
        if any(kw in line_lower for kw in exp_keywords) and len(line.strip()) < 40:
            in_exp = True
            continue
        if in_exp and any(kw in line_lower for kw in stop_keywords) and len(line.strip()) < 40:
            break
        if in_exp:
            exp_section += line + '\n'

    # Parse job entries from experience section
    date_pattern = r'(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s*\d{2,4}|' \
                   r'\d{1,2}/\d{2,4}|\d{4}\s*[-–]\s*(?:\d{4}|present|current|now)'

    chunks = re.split(date_pattern, exp_section, flags=re.IGNORECASE)
    dates = re.findall(date_pattern, exp_section, flags=re.IGNORECASE)

    for i, chunk in enumerate(chunks[:10]):
        chunk = chunk.strip()
        if len(chunk) > 20:
            lines_c = [l.strip() for l in chunk.split('\n') if l.strip()]
            entry = {
                "title": lines_c[0] if lines_c else "",
                "company": lines_c[1] if len(lines_c) > 1 else "",
                "dates": dates[i - 1] if i > 0 and i - 1 < len(dates) else "",
                "description": ' '.join(lines_c[2:4]) if len(lines_c) > 2 else "",
            }
            if entry["title"]:
                experience.append(entry)

    return experience[:8]


def extract_education(text: str) -> list:
    education = []
    lines = text.split('\n')
    edu_keywords = ['bachelor', 'master', 'mba', 'phd', 'b.sc', 'm.sc', 'b.eng', 'diploma',
                    'degree', 'university', 'college', 'institute', 'school', 'faculty']
    for line in lines:
        line_lower = line.lower()
        if any(kw in line_lower for kw in edu_keywords) and len(line.strip()) > 5:
            education.append({"text": line.strip()})
    return education[:5]


def extract_languages(text: str) -> list:
    lang_keywords = ['arabic', 'english', 'french', 'hindi', 'urdu', 'german', 'chinese',
                     'spanish', 'italian', 'portuguese', 'russian', 'japanese', 'korean']
    found = []
    text_lower = text.lower()
    for lang in lang_keywords:
        if lang in text_lower:
            found.append(lang.capitalize())
    return found


def extract_summary(text: str) -> str:
    lines = text.split('\n')
    summary_keywords = ['summary', 'objective', 'profile', 'about', 'overview']
    for i, line in enumerate(lines):
        if any(kw in line.lower() for kw in summary_keywords) and len(line.strip()) < 40:
            summary_lines = []
            for j in range(i + 1, min(i + 8, len(lines))):
                l = lines[j].strip()
                if l and len(l) > 10:
                    summary_lines.append(l)
                elif not l and summary_lines:
                    break
            if summary_lines:
                return ' '.join(summary_lines)
    # Fallback: first paragraph
    paragraphs = [p.strip() for p in text.split('\n\n') if len(p.strip()) > 50]
    return paragraphs[0][:500] if paragraphs else ""


def parse_resume(file_path: str, filename: str) -> dict:
    raw_text = extract_text(file_path)
    return {
        "filename": filename,
        "file_path": file_path,
        "raw_text": raw_text,
        "name": extract_name(raw_text),
        "email": extract_email(raw_text),
        "phone": extract_phone(raw_text),
        "location": None,
        "summary": extract_summary(raw_text),
        "skills": extract_skills(raw_text),
        "experience": extract_experience(raw_text),
        "education": extract_education(raw_text),
        "languages": extract_languages(raw_text),
    }
